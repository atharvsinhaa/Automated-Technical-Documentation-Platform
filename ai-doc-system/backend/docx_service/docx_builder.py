"""
docx_service/docx_builder.py
────────────────────────────────────────────────────────────────
Core DOCX Construction Engine.

Builds professionally formatted Word documents with:
  - Cover page
  - Table of Contents
  - Revision history
  - Enterprise-styled headings, tables, lists
  - Embedded Mermaid diagrams (rendered to PNG)
  - Headers & footers with page numbers
  - Section numbering

Uses python-docx for document assembly and the MermaidToPng
renderer for diagram embedding.
"""

from __future__ import annotations

import os
import re
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from .styles import (
    apply_enterprise_styles,
    format_enterprise_table,
    NAVY, ACCENT_BLUE, DARK_GRAY, MID_GRAY, WHITE, LIGHT_GRAY,
    FONT_HEADING, FONT_BODY, FONT_CODE,
    CODE_BG, CODE_FG,
    COVER_ACCENT,
)
from .markdown_parser import (
    parse_markdown,
    parse_inline,
    BlockType,
    ContentBlock,
    InlineSegment,
)
from .mermaid_renderer import MermaidToPng

def get_png_dimensions(filepath: str) -> Tuple[int, int]:
    """Securely extract width and height from PNG header without heavy dependencies."""
    try:
        with open(filepath, 'rb') as f:
            data = f.read(24)
            if data[:8] == b'\x89PNG\r\n\x1a\n':
                w = int.from_bytes(data[16:20], byteorder='big')
                h = int.from_bytes(data[20:24], byteorder='big')
                return w, h
    except Exception:
        pass
    return 0, 0


class DocxBuilder:
    """
    Enterprise DOCX document builder.

    Provides a fluent API for assembling professional Word
    documents with all enterprise features.
    """

    def __init__(
        self,
        project_name: str = "AI Documentation",
        repo_name: str = "",
        version: str = "1.0",
        author: str = "AI Documentation Platform",
        verbose: bool = True,
    ):
        self.project_name = project_name
        self.repo_name = repo_name
        self.version = version
        self.author = author
        self.verbose = verbose

        self.doc = Document()
        self._mermaid = MermaidToPng(verbose=verbose)
        self._temp_dir = tempfile.mkdtemp(prefix="docx_diagrams_")
        self._diagram_counter = 0

        # Apply enterprise styles
        apply_enterprise_styles(self.doc)

    # ══════════════════════════════════════════════════════════
    #  COVER PAGE
    # ══════════════════════════════════════════════════════════

    def add_cover_page(
        self,
        title: str,
        subtitle: str = "Architecture Documentation",
        project: str = "",
        repo: str = "",
        version: str = "",
        date: str = "",
        page_break: bool = True,
    ) -> "DocxBuilder":
        """
        Add a professional cover page.
        """
        project = project or self.project_name
        repo = repo or self.repo_name
        version = version or self.version
        date = date or datetime.now().strftime("%B %d, %Y")

        # Top spacing
        for _ in range(4):
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)

        # Accent line
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("━" * 40)
        run.font.color.rgb = ACCENT_BLUE
        run.font.size = Pt(14)

        # Title
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(title)
        run.font.size = Pt(36)
        run.font.color.rgb = NAVY
        run.font.name = FONT_HEADING
        run.bold = True

        # Subtitle
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(20)
        run = p.add_run(subtitle)
        run.font.size = Pt(16)
        run.font.color.rgb = MID_GRAY
        run.font.name = FONT_HEADING

        # Accent line
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("━" * 40)
        run.font.color.rgb = ACCENT_BLUE
        run.font.size = Pt(14)

        # Spacing
        for _ in range(3):
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)

        # Metadata table
        meta = [
            ("Project", project),
            ("Repository", repo or "—"),
            ("Version", version),
            ("Generated", date),
            ("Author", self.author),
        ]

        for label, value in meta:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.space_before = Pt(2)

            run = p.add_run(f"{label}: ")
            run.font.name = FONT_BODY
            run.font.size = Pt(11)
            run.font.color.rgb = MID_GRAY

            run = p.add_run(value)
            run.font.name = FONT_BODY
            run.font.size = Pt(11)
            run.font.color.rgb = DARK_GRAY
            run.bold = True

        # Page break after cover
        if page_break:
            self.doc.add_page_break()
        else:
            self.doc.add_paragraph()  # just spacing

        return self

    # ══════════════════════════════════════════════════════════
    #  TABLE OF CONTENTS
    # ══════════════════════════════════════════════════════════

    def add_toc(self, title: str = "Table of Contents", page_break: bool = True) -> "DocxBuilder":
        """
        Add a Table of Contents that auto-updates in Word.

        Uses Word's built-in TOC field code. The TOC will
        populate when the user opens the document and presses
        Ctrl+A → F9, or when Print Preview is triggered.
        """
        # TOC heading
        heading = self.doc.add_heading(title, level=1)

        # Add TOC field
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run()

        fld_char_begin = parse_xml(
            f'<w:fldChar {nsdecls("w")} w:fldCharType="begin" w:dirty="true"/>'
        )
        run._r.append(fld_char_begin)

        instr_text = parse_xml(
            f'<w:instrText {nsdecls("w")} xml:space="preserve">'
            f' TOC \\o "1-3" \\h \\z \\u '
            f'</w:instrText>'
        )
        run._r.append(instr_text)

        fld_char_separate = parse_xml(
            f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>'
        )
        run._r.append(fld_char_separate)

        fld_char_end = parse_xml(
            f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'
        )
        run2 = paragraph.add_run()
        run2._r.append(fld_char_end)

        # Ensure updateFields is enabled in doc settings to auto-update TOC
        element = parse_xml(f'<w:updateFields {nsdecls("w")} w:val="true"/>')
        self.doc.settings.element.append(element)

        # Page break after TOC
        if page_break:
            self.doc.add_page_break()
        else:
            self.doc.add_paragraph()

        return self

    # ══════════════════════════════════════════════════════════
    #  REVISION HISTORY
    # ══════════════════════════════════════════════════════════

    def add_revision_history(
        self,
        entries: Optional[List[Tuple[str, str, str, str]]] = None,
    ) -> "DocxBuilder":
        """
        Add a revision history table.

        Args:
            entries: List of (version, date, author, description) tuples.
                     If None, adds a single entry for the current version.
        """
        self.doc.add_heading("Revision History", level=1)

        if entries is None:
            entries = [(
                self.version,
                datetime.now().strftime("%Y-%m-%d"),
                self.author,
                "Initial document generation",
            )]

        headers = ["Version", "Date", "Author", "Description"]
        table = self.doc.add_table(
            rows=1 + len(entries),
            cols=4,
        )

        # Header row
        for j, header in enumerate(headers):
            cell = table.rows[0].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(header)
            run.bold = True

        # Data rows
        for i, entry in enumerate(entries):
            for j, value in enumerate(entry):
                cell = table.rows[i + 1].cells[j]
                cell.text = value

        format_enterprise_table(table)

        self.doc.add_paragraph()  # spacing

        return self

    # ══════════════════════════════════════════════════════════
    #  CONTENT PRIMITIVES
    # ══════════════════════════════════════════════════════════

    def add_heading(
        self,
        text: str,
        level: int = 1,
    ) -> "DocxBuilder":
        """Add a styled heading."""
        level = max(1, min(level, 4))
        self.doc.add_heading(text, level=level)
        return self

    def add_paragraph(
        self,
        text: str,
        style: str = "Normal",
        bold: bool = False,
        italic: bool = False,
    ) -> "DocxBuilder":
        """
        Add a paragraph with optional inline formatting.
        Parses **bold**, *italic*, and `code` automatically.
        """
        p = self.doc.add_paragraph(style=style)

        segments = parse_inline(text)
        for seg in segments:
            run = p.add_run(seg.text)
            run.bold = seg.bold or bold
            run.italic = seg.italic or italic
            if seg.code:
                run.font.name = FONT_CODE
                run.font.size = Pt(9)
                run.font.color.rgb = CODE_FG

        return self

    def add_table(
        self,
        headers: List[str],
        rows: List[List[str]],
    ) -> "DocxBuilder":
        """Add a professionally formatted table."""
        if not headers and not rows:
            return self

        num_cols = len(headers) if headers else len(rows[0])
        num_rows = (1 if headers else 0) + len(rows)

        table = self.doc.add_table(rows=num_rows, cols=num_cols)

        row_idx = 0

        # Header row
        if headers:
            for j, header in enumerate(headers):
                cell = table.rows[0].cells[j]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(header)
                run.bold = True
            row_idx = 1

        # Data rows
        for row_data in rows:
            for j, value in enumerate(row_data):
                if j < num_cols:
                    cell = table.rows[row_idx].cells[j]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    # Parse inline formatting in cells
                    segments = parse_inline(value)
                    for seg in segments:
                        run = p.add_run(seg.text)
                        run.bold = seg.bold
                        run.italic = seg.italic
                        if seg.code:
                            run.font.name = FONT_CODE
                            run.font.size = Pt(9)
            row_idx += 1

        format_enterprise_table(table)

        # Spacing after table
        self.doc.add_paragraph()

        return self

    def add_bullet_list(
        self,
        items: List[str],
    ) -> "DocxBuilder":
        """Add a bullet list with inline formatting."""
        for item in items:
            p = self.doc.add_paragraph(style="EnterpriseBullet")
            segments = parse_inline(item)
            for seg in segments:
                run = p.add_run(seg.text)
                run.bold = seg.bold
                run.italic = seg.italic
                
                # Apply standard font styling directly to the run so we don't overwrite 
                # the paragraph's native bullet character font
                if seg.code:
                    run.font.name = FONT_CODE
                    run.font.size = Pt(9)
                    run.font.color.rgb = CODE_FG
                else:
                    run.font.name = FONT_BODY
                    run.font.size = Pt(11)
                    run.font.color.rgb = DARK_GRAY
        return self

    def add_numbered_list(
        self,
        items: List[str],
    ) -> "DocxBuilder":
        """Add a numbered list with inline formatting."""
        for item in items:
            p = self.doc.add_paragraph(style="EnterpriseNumber")
            segments = parse_inline(item)
            for seg in segments:
                run = p.add_run(seg.text)
                run.bold = seg.bold
                run.italic = seg.italic
                
                if seg.code:
                    run.font.name = FONT_CODE
                    run.font.size = Pt(9)
                    run.font.color.rgb = CODE_FG
                else:
                    run.font.name = FONT_BODY
                    run.font.size = Pt(11)
                    run.font.color.rgb = DARK_GRAY
        return self

    def add_code_block(
        self,
        code: str,
        language: str = "",
    ) -> "DocxBuilder":
        """
        Add a code block with monospace font and gray background.
        """
        for line in code.split("\n"):
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

            run = p.add_run(line or " ")
            run.font.name = FONT_CODE
            run.font.size = Pt(9)
            run.font.color.rgb = CODE_FG

            # Background shading
            pPr = p._p.get_or_add_pPr()
            shading = parse_xml(
                f'<w:shd {nsdecls("w")} '
                f'w:fill="{CODE_BG}" w:val="clear"/>'
            )
            pPr.append(shading)

        # Spacing after code block
        self.doc.add_paragraph()

        return self

    def add_mermaid_diagram(
        self,
        mermaid_code: str,
        caption: str = "",
    ) -> "DocxBuilder":
        """
        Render a Mermaid diagram to PNG and embed it.
        """
        self._diagram_counter += 1
        name = f"diagram_{self._diagram_counter}"

        png_path = self._mermaid.render(
            mermaid_code=mermaid_code,
            output_dir=self._temp_dir,
            name=name,
        )

        if png_path and os.path.exists(png_path):
            self.add_image(
                path=png_path,
                caption=caption or f"Figure {self._diagram_counter}",
                width_inches=6.0,
            )
        else:
            self._log("  [docx] ⚠ Mermaid render failed — structured text fallback")
            if caption:
                p = self.doc.add_paragraph()
                run = p.add_run(f"◈ {caption}")
                run.bold = True
                run.font.name = FONT_HEADING
                run.font.size = Pt(11)
                run.font.color.rgb = ACCENT_BLUE
                p.paragraph_format.space_before = Pt(6)

            p = self.doc.add_paragraph()
            note_run = p.add_run(
                "⚠  Diagram rendering requires Mermaid CLI. "
                "Install: npm install -g @mermaid-js/mermaid-cli"
            )
            note_run.italic = True
            note_run.font.size = Pt(9)
            note_run.font.color.rgb = MID_GRAY
            p.paragraph_format.space_after = Pt(6)

            if mermaid_code:
                self._render_mermaid_as_text_fallback(mermaid_code)

        return self

    def _render_mermaid_as_text_fallback(self, mermaid_code: str):
        """Extract meaningful info from Mermaid code when PNG render fails."""
        import re
        lines = mermaid_code.strip().split("\n")
        if not lines: return
        dtype = lines[0].strip().lower()

        if "sequencediagram" in dtype:
            interactions = [l.strip() for l in lines[1:]
                            if "->>" in l or "-->" in l and len(l.strip()) > 3]
            if interactions:
                self.add_paragraph("Interaction Steps:", bold=True)
                self.add_numbered_list(interactions[:10])

        elif "classdiagram" in dtype:
            class_names = [m.group(1) for l in lines[1:]
                           for m in [re.match(r'\s*class\s+(\w+)', l.strip())] if m]
            if class_names:
                self.add_paragraph(f"Classes: {', '.join(f'`{c}`' for c in class_names[:8])}")

        elif "flowchart" in dtype or "graph" in dtype:
            nodes = re.findall(r'\w+\["([^"]+)"\]', mermaid_code)
            if nodes:
                self.add_paragraph("Components:", bold=True)
                self.add_bullet_list(nodes[:8])

        elif "erdiagram" in dtype:
            entities = re.findall(r'^\s*(\w+)\s*\{', mermaid_code, re.MULTILINE)
            if entities:
                self.add_paragraph(f"Entities: {', '.join(f'`{e}`' for e in entities[:8])}")

    def add_image(
        self,
        path: str,
        caption: str = "",
        width_inches: float = 6.0,
    ) -> "DocxBuilder":
        """
        Embed an image with an optional caption.
        """
        if not os.path.exists(path):
            self._log(f"  [docx] ⚠ Image not found: {path}")
            return self

        # Image auto-scaling based on aspect ratio
        w_px, h_px = get_png_dimensions(path)
        if w_px > 0 and h_px > 0:
            aspect_ratio = h_px / w_px
            if aspect_ratio > 1.2:  # It's a tall diagram
                width_inches = 4.5
            else:
                width_inches = 6.0

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(width_inches))

        # Caption
        if caption:
            cap_p = self.doc.add_paragraph()
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_p.paragraph_format.space_before = Pt(4)
            cap_p.paragraph_format.space_after = Pt(12)
            run = cap_p.add_run(caption)
            run.font.size = Pt(9)
            run.font.color.rgb = MID_GRAY
            run.italic = True

        return self

    def add_section_break(self) -> "DocxBuilder":
        """Add a page break."""
        self.doc.add_page_break()
        return self

    # ══════════════════════════════════════════════════════════
    #  HEADERS & FOOTERS
    # ══════════════════════════════════════════════════════════

    def setup_headers_footers(
        self,
        title: str = "",
        version: str = "",
    ) -> "DocxBuilder":
        """
        Configure document headers and footers.

        Header: Document title (left) | Version (right)
        Footer: "Confidential" (left) | Page X of Y (right)
        """
        title = title or self.project_name
        version = version or f"v{self.version}"

        for section in self.doc.sections:
            # ── Header ────────────────────────────────────────
            header = section.header
            header.is_linked_to_previous = False

            # Clear existing
            for p in header.paragraphs:
                p.clear()

            h_para = header.paragraphs[0]

            # Left: title
            run = h_para.add_run(title)
            run.font.name = FONT_BODY
            run.font.size = Pt(8)
            run.font.color.rgb = MID_GRAY

            # Tab to right
            run = h_para.add_run("\t\t")

            # Right: version
            run = h_para.add_run(version)
            run.font.name = FONT_BODY
            run.font.size = Pt(8)
            run.font.color.rgb = MID_GRAY

            # ── Footer ────────────────────────────────────────
            footer = section.footer
            footer.is_linked_to_previous = False

            # Clear existing
            for p in footer.paragraphs:
                p.clear()

            f_para = footer.paragraphs[0]

            # Left: confidential
            run = f_para.add_run("Confidential")
            run.font.name = FONT_BODY
            run.font.size = Pt(8)
            run.font.color.rgb = MID_GRAY
            run.italic = True

            # Tab to right
            run = f_para.add_run("\t\t")

            # Right: Page X of Y
            run = f_para.add_run("Page ")
            run.font.name = FONT_BODY
            run.font.size = Pt(8)
            run.font.color.rgb = MID_GRAY

            # Page number field
            fld_xml = (
                f'<w:fldSimple {nsdecls("w")} w:instr=" PAGE \\* MERGEFORMAT ">'
                f'<w:r><w:t>1</w:t></w:r>'
                f'</w:fldSimple>'
            )
            fld = parse_xml(fld_xml)
            f_para._p.append(fld)

            run = f_para.add_run(" of ")
            run.font.name = FONT_BODY
            run.font.size = Pt(8)
            run.font.color.rgb = MID_GRAY

            # Total pages field
            fld_xml2 = (
                f'<w:fldSimple {nsdecls("w")} w:instr=" NUMPAGES \\* MERGEFORMAT ">'
                f'<w:r><w:t>1</w:t></w:r>'
                f'</w:fldSimple>'
            )
            fld2 = parse_xml(fld_xml2)
            f_para._p.append(fld2)

            # Set margins
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        return self

    # ══════════════════════════════════════════════════════════
    #  MARKDOWN → DOCX (FULL CONVERTER)
    # ══════════════════════════════════════════════════════════

    def from_markdown(self, md_text: str) -> "DocxBuilder":
        """
        Parse markdown and add all content to the document.
        Mermaid diagrams are rendered to PNG and embedded.
        Tables are converted to Word tables.
        """
        blocks = parse_markdown(md_text)

        for block in blocks:
            if block.block_type == BlockType.HEADING:
                if block.level == 1:
                    continue  # Skip top-level title, it's handled by the cover page
                # Map Markdown ## -> Heading 1, ### -> Heading 2
                adjusted_level = max(1, block.level - 1)
                self.add_heading(block.text, adjusted_level)

            elif block.block_type == BlockType.PARAGRAPH:
                self.add_paragraph(block.text)

            elif block.block_type == BlockType.TABLE:
                self.add_table(block.headers, block.rows)

            elif block.block_type == BlockType.BULLET_LIST:
                self.add_bullet_list(block.items)

            elif block.block_type == BlockType.NUMBERED_LIST:
                self.add_numbered_list(block.items)

            elif block.block_type == BlockType.MERMAID_DIAGRAM:
                self.add_mermaid_diagram(block.code)

            elif block.block_type == BlockType.CODE_BLOCK:
                self.add_code_block(block.code, block.language)

            elif block.block_type == BlockType.HORIZONTAL_RULE:
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run("─" * 50)
                run.font.color.rgb = LIGHT_GRAY
                run.font.size = Pt(8)

        return self

    # ══════════════════════════════════════════════════════════
    #  SAVE
    # ══════════════════════════════════════════════════════════

    def save(self, output_path: str) -> str:
        """
        Save the document to disk.
        Returns the output path.
        """
        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)

        self.doc.save(output_path)

        size_kb = os.path.getsize(output_path) / 1024
        self._log(f"[docx] ✓ Saved: {output_path} ({size_kb:.1f} KB)")

        return output_path

    def _log(self, msg: str):
        if self.verbose:
            print(msg, flush=True)
